"""
Professional DOCX Summary Creator
Creates beautifully formatted Word documents from book summaries.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============================================================
# STYLE CONFIGURATION
# ============================================================

COLORS = {
    "primary": RGBColor(0x1A, 0x56, 0xDB),     # Deep blue
    "secondary": RGBColor(0x6B, 0x72, 0x80),    # Gray
    "accent": RGBColor(0xE8, 0x4D, 0x3D),       # Warm red
    "dark": RGBColor(0x1F, 0x29, 0x37),          # Near black
    "light_gray": RGBColor(0xF3, 0xF4, 0xF6),   # Light bg
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "text": RGBColor(0x37, 0x41, 0x51),          # Body text
    "quote_bg": RGBColor(0xEF, 0xF6, 0xFF),      # Light blue
}

FONTS = {
    "heading": "Calibri",
    "body": "Georgia",
    "quote": "Georgia",
    "mono": "Consolas",
}


# ============================================================
# DOCUMENT BUILDER
# ============================================================

class DocxSummaryCreator:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    def _setup_page(self):
        """Configure page layout."""
        section = self.doc.sections[0]
        section.page_width = Cm(21)    # A4
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles

        # Normal style
        normal = style['Normal']
        normal.font.name = FONTS["body"]
        normal.font.size = Pt(11)
        normal.font.color.rgb = COLORS["text"]
        normal.paragraph_format.line_spacing = 1.4
        normal.paragraph_format.space_after = Pt(6)

        # Title style
        title_style = style['Title']
        title_style.font.name = FONTS["heading"]
        title_style.font.size = Pt(28)
        title_style.font.color.rgb = COLORS["primary"]
        title_style.font.bold = True
        title_style.paragraph_format.space_after = Pt(4)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Heading 1
        h1 = style['Heading 1']
        h1.font.name = FONTS["heading"]
        h1.font.size = Pt(20)
        h1.font.color.rgb = COLORS["primary"]
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)

        # Heading 2
        h2 = style['Heading 2']
        h2.font.name = FONTS["heading"]
        h2.font.size = Pt(16)
        h2.font.color.rgb = COLORS["dark"]
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)

        # Heading 3
        h3 = style['Heading 3']
        h3.font.name = FONTS["heading"]
        h3.font.size = Pt(13)
        h3.font.color.rgb = COLORS["dark"]
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)

        # Create Quote style
        try:
            quote_style = style.add_style('BookQuote', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            quote_style = style['BookQuote']
        quote_style.font.name = FONTS["quote"]
        quote_style.font.size = Pt(11)
        quote_style.font.italic = True
        quote_style.font.color.rgb = COLORS["secondary"]
        quote_style.paragraph_format.left_indent = Cm(1.5)
        quote_style.paragraph_format.right_indent = Cm(1.5)
        quote_style.paragraph_format.space_before = Pt(8)
        quote_style.paragraph_format.space_after = Pt(8)

    def add_cover_page(self, title, author, pages, file_size):
        """Create a professional cover page."""
        # Add spacing before title
        for _ in range(4):
            self.doc.add_paragraph("")

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("📖 COGNOSCERE REVIEW")
        run.font.name = FONTS["heading"]
        run.font.size = Pt(14)
        run.font.color.rgb = COLORS["secondary"]
        run.font.bold = True

        # Book title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = FONTS["heading"]
        run.font.size = Pt(32)
        run.font.color.rgb = COLORS["primary"]
        run.font.bold = True

        # Horizontal line
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = COLORS["primary"]
        run.font.size = Pt(10)

        # Author
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"oleh {author}")
        run.font.name = FONTS["heading"]
        run.font.size = Pt(18)
        run.font.color.rgb = COLORS["dark"]
        run.font.italic = True

        # Spacer
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        # Info box
        info_table = self.doc.add_table(rows=1, cols=3)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cells = info_table.rows[0].cells
        info_items = [
            ("📄", f"{pages} halaman"),
            ("📏", f"{file_size} MB"),
            ("🤖", "AI Summary"),
        ]

        for i, (icon, text) in enumerate(info_items):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{icon} {text}")
            run.font.name = FONTS["heading"]
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS["secondary"]

        # Add spacing
        for _ in range(6):
            self.doc.add_paragraph("")

        # Footer note
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("COGNOSCERE Comprehensive Knowledge Extraction")
        run.font.name = FONTS["heading"]
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS["secondary"]
        run.font.italic = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("COGNOSCERE v2.0 • Powered by Kimi AI")
        run.font.name = FONTS["heading"]
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS["secondary"]

        # Page break
        self.doc.add_page_break()

    def add_table_of_contents_placeholder(self):
        """Add a simple table of contents."""
        self.doc.add_heading("Daftar Isi", level=1)

        toc_items = [
            "1. Ikhtisar Buku (Overview)",
            "2. Tesis Utama (Core Thesis)",
            "3. Rangkuman Per Bab / Bagian",
            "4. Konsep & Ide Kunci",
            "5. Kutipan Penting",
            "6. Pelajaran Utama (Key Takeaways)",
            "7. Penerapan Praktis",
            "8. Kritik & Perspektif",
            "9. Kesimpulan",
        ]

        for item in toc_items:
            p = self.doc.add_paragraph()
            run = p.add_run(item)
            run.font.name = FONTS["heading"]
            run.font.size = Pt(12)
            run.font.color.rgb = COLORS["dark"]
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.5)

        self.doc.add_page_break()

    def _parse_markdown_to_docx(self, markdown_text):
        """Convert markdown-formatted summary to DOCX elements."""
        lines = markdown_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # Heading 1: ### or ##
            if line.startswith('### '):
                heading_text = re.sub(r'^###\s+\d*\.?\s*', '', line).strip()
                self.doc.add_heading(heading_text, level=1)
                i += 1
                continue

            if line.startswith('## '):
                heading_text = re.sub(r'^##\s+\d*\.?\s*', '', line).strip()
                self.doc.add_heading(heading_text, level=1)
                i += 1
                continue

            if line.startswith('# '):
                heading_text = re.sub(r'^#\s+\d*\.?\s*', '', line).strip()
                self.doc.add_heading(heading_text, level=1)
                i += 1
                continue

            # Sub-heading: #### or bold line
            if line.startswith('#### '):
                heading_text = re.sub(r'^####\s+', '', line).strip()
                self.doc.add_heading(heading_text, level=2)
                i += 1
                continue

            # Horizontal rule
            if line.startswith('---') or line.startswith('___') or line.startswith('***'):
                p = self.doc.add_paragraph()
                run = p.add_run("━" * 60)
                run.font.color.rgb = COLORS["light_gray"]
                run.font.size = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue

            # Quote block (starts with > or ")
            if line.startswith('>') or (line.startswith('"') and line.endswith('"')):
                quote_text = line.lstrip('> ').strip('"').strip()
                # Collect multi-line quotes
                while i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                    i += 1
                    quote_text += ' ' + lines[i].strip().lstrip('> ')

                try:
                    p = self.doc.add_paragraph(style='BookQuote')
                except KeyError:
                    p = self.doc.add_paragraph()
                run = p.add_run(f'"{quote_text}"')
                run.font.italic = True
                run.font.color.rgb = COLORS["secondary"]
                i += 1
                continue

            # Bullet point
            if line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
                bullet_text = line[2:].strip()
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, bullet_text)
                i += 1
                continue

            # Numbered list
            num_match = re.match(r'^(\d+)\.\s+(.+)', line)
            if num_match:
                list_text = num_match.group(2).strip()
                p = self.doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, list_text)
                i += 1
                continue

            # Regular paragraph - collect consecutive lines
            para_lines = [line]
            while i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if (not next_line or
                    next_line.startswith('#') or
                    next_line.startswith('- ') or
                    next_line.startswith('* ') or
                    next_line.startswith('• ') or
                    next_line.startswith('>') or
                    next_line.startswith('---') or
                    re.match(r'^\d+\.\s', next_line)):
                    break
                para_lines.append(next_line)
                i += 1

            full_text = ' '.join(para_lines)
            p = self.doc.add_paragraph()
            self._add_formatted_text(p, full_text)
            i += 1

    def _add_formatted_text(self, paragraph, text):
        """Add text with inline formatting (bold, italic) to a paragraph."""
        # Process bold and italic markers
        parts = re.split(r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)', text)

        for part in parts:
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                run.font.color.rgb = COLORS["dark"]
            elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)

    def add_summary_content(self, markdown_summary):
        """Add the full summary content from markdown."""
        self._parse_markdown_to_docx(markdown_summary)

    def add_footer(self):
        """Add a footer with generation info."""
        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = COLORS["light_gray"]
        run.font.size = Pt(8)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\nCOGNOSCERE v2.0 — Comprehensive Knowledge Extraction System\n")
        run.font.name = FONTS["heading"]
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS["secondary"]
        run.font.italic = True

        run = p.add_run("Powered by Kimi (Moonshot AI)\n")
        run.font.name = FONTS["heading"]
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS["secondary"]

        run = p.add_run("⚠️ AI extraction mungkin tidak 100% akurat. Untuk pemahaman mendalam, baca buku aslinya.")
        run.font.name = FONTS["heading"]
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["accent"]
        run.font.italic = True

    def save(self, filepath):
        """Save the DOCX file."""
        os.makedirs(os.path.dirname(filepath) or '.', exist_ok=True)
        self.doc.save(filepath)
        print(f"  📄 DOCX saved: {filepath}")


def create_book_summary_docx(metadata, summary_text, output_path):
    """
    Create a complete, professionally formatted book summary DOCX.

    Args:
        metadata: dict with title, author, pages, file_size_mb
        summary_text: Markdown-formatted summary text
        output_path: Where to save the DOCX
    """
    creator = DocxSummaryCreator()

    # Cover page
    creator.add_cover_page(
        title=metadata.get("title", "Unknown"),
        author=metadata.get("author", "Unknown"),
        pages=metadata.get("pages", 0),
        file_size=metadata.get("file_size_mb", 0)
    )

    # Table of contents
    creator.add_table_of_contents_placeholder()

    # Main content
    creator.add_summary_content(summary_text)

    # Footer
    creator.add_footer()

    # Save
    creator.save(output_path)


if __name__ == "__main__":
    # Test with sample content
    sample_metadata = {
        "title": "Atomic Habits",
        "author": "James Clear",
        "pages": 320,
        "file_size_mb": 2.5
    }

    sample_summary = """### 1. IKHTISAR BUKU (Overview)

**Atomic Habits** adalah buku yang membahas tentang bagaimana perubahan kecil dapat menghasilkan hasil yang luar biasa. James Clear menjelaskan bahwa kebiasaan adalah bunga majemuk dari perbaikan diri.

Buku ini ditulis berdasarkan pengalaman pribadi penulis dan penelitian ilmiah tentang pembentukan kebiasaan.

### 2. TESIS UTAMA (Core Thesis)

Argumen utama Clear adalah bahwa **perubahan sebesar 1% setiap hari** akan menghasilkan perbaikan 37 kali lipat dalam setahun. Ini bukan tentang membuat perubahan besar dan dramatis, tetapi tentang konsistensi dalam hal-hal kecil.

### 3. KONSEP & IDE KUNCI

- **The 1% Rule**: Perbaikan kecil yang konsisten menghasilkan hasil besar seiring waktu
- **Identity-Based Habits**: Fokus pada siapa Anda ingin menjadi, bukan apa yang ingin dicapai
- **The Four Laws of Behavior Change**: Make it obvious, attractive, easy, and satisfying

### 4. KUTIPAN PENTING

> "You do not rise to the level of your goals. You fall to the level of your systems."
> — Terjemahan: Anda tidak naik ke level tujuan Anda. Anda jatuh ke level sistem Anda.

### 5. PELAJARAN UTAMA (Key Takeaways)

1. Fokus pada sistem, bukan tujuan
2. Identitas mendorong perilaku
3. Lingkungan lebih penting dari motivasi
4. Kebiasaan dimulai dari yang sangat kecil

### 6. KESIMPULAN

Atomic Habits adalah panduan praktis yang mengubah cara kita berpikir tentang kebiasaan dan perubahan diri."""

    output = "/sessions/ecstatic-sleepy-johnson/mnt/Downloads/BookSummarizer/test_summary.docx"
    create_book_summary_docx(sample_metadata, sample_summary, output)
    print(f"✅ Test DOCX created: {output}")
